from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
FIG = REPORT / "figures"
OUT = REPORT / "final_report.docx"


BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(89, 89, 89)


def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 15, BLUE),
        ("Heading 2", 12.5, BLUE),
        ("Heading 3", 11.5, BLUE),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("SmartFlow: Adaptive Traffic Signal Control with Reinforcement Learning")
    set_font(run, size=18, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Zhonghuan Tang")
    set_font(run, size=11, color=GRAY)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_font(run, size=9.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "999999")


def set_table_widths(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_appendix_table(doc, title, headers, rows, widths):
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    set_table_widths(table, widths)
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], "EEEEEE")
        set_cell_text(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], value, align=align)
    set_table_widths(table, widths)


def add_appendix_image(doc, title, image_name, width=5.8):
    add_heading(doc, title, level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(str(FIG / image_name), width=Inches(width))


def build():
    doc = Document()
    setup_document(doc)
    part1 = pd.read_csv(ROOT / "experiments" / "simple_intersection" / "presentation_results.csv")
    part2 = pd.read_csv(ROOT / "experiments" / "complex_network" / "results_complex.csv")

    add_title(doc)

    add_heading(doc, "Introduction")
    add_body(
        doc,
        "Urban traffic control is a useful setting for studying sequential decision-making because signal timing affects both the vehicles currently waiting and the congestion that will appear later. A fixed-cycle traffic signal follows the same schedule even when demand changes, which means it may waste green time on an empty lane or fail to respond when one direction becomes crowded. Adaptive control is meant to solve this problem by changing signal behavior in response to observed conditions. Reinforcement learning is a natural approach because it lets a controller improve through interaction with a traffic environment rather than relying only on a rule written in advance.",
    )
    add_body(
        doc,
        "The SmartFlow project began with a simple single-intersection experiment. The goal of that first experiment was to test whether reinforcement learning could learn a useful traffic signal policy and how its performance compared with common baselines. The project then extended the same idea to a more complex two-intersection corridor. This extension is important because real traffic networks are rarely isolated. A decision at one signal can affect queues at another signal several steps later. That delayed interaction is exactly where a long-term learning method should become more valuable than a purely reactive heuristic.",
    )

    add_heading(doc, "Part I: Original Single-Intersection Experiment")
    add_body(
        doc,
        "The original experiment defined the intersection as a Markov Decision Process. The state was S = (Q_NS, Q_EW, Phase), where Q_NS is the north-south queue, Q_EW is the east-west queue, and Phase indicates the current green-light direction. To make tabular learning manageable, each queue was placed into one of three buckets: low traffic for 0 to 3 vehicles, medium traffic for 4 to 8 vehicles, and high traffic for 9 or more vehicles. This representation reduces the state space while still giving the agent basic information about which direction is more congested.",
    )
    add_body(
        doc,
        "The action space contained two choices: Stay or Switch. Stay means the signal keeps the current green phase, while Switch means the signal changes to the other direction. The reward function was R = -(Q_NS + Q_EW) - C, where the first term penalizes waiting vehicles and the switching cost C discourages the controller from changing phases too frequently. This design captures the main tradeoff in signal control. The agent should reduce queues, but it should not create unrealistic or unsafe oscillation by switching at every time step.",
    )
    add_body(
        doc,
        "The algorithms compared in the original presentation were Q-Learning, SARSA, Monte Carlo learning, static fixed-cycle timing, a greedy longest-queue heuristic, and Value Iteration as an optimal reference. These methods create a useful range of comparisons. Static timing represents a simple non-adaptive baseline. Greedy timing represents a strong local heuristic because it directly serves the longer current queue. Value Iteration represents the best policy available under the simplified model assumptions. The reinforcement learning methods show whether an agent can learn adaptive behavior from repeated simulated experience.",
    )
    add_body(
        doc,
        "The preserved presentation results show that SARSA slightly outperformed static timing, while the greedy heuristic and Value Iteration performed better than the learned tabular agents. This result should be interpreted carefully. SARSA beating static timing supports the basic claim that reinforcement learning can learn an adaptive signal policy. However, greedy performing strongly does not mean reinforcement learning is useless. In a single two-phase intersection, the current local queue is almost the whole problem. A longest-queue rule has direct access to that information, so it can be very effective without learning any long-term strategy.",
    )

    add_heading(doc, "Why the Simple Experiment Is Limited")
    add_body(
        doc,
        "The main limitation of the first experiment is that the environment is too local. There are no neighboring intersections, no vehicles traveling between signals, and no downstream spillback. The greedy heuristic only needs to ask which queue is longer right now. Because there is no network effect, that short-term decision often lines up with good performance. The reinforcement learning agents also used coarse buckets instead of exact queue counts, so they had less detailed state information than the greedy rule. This creates a perception gap: the heuristic appears stronger partly because it sees exact local queues in a very simple environment.",
    )
    add_body(
        doc,
        "This limitation motivated the second experiment. A stronger test of reinforcement learning should include delayed consequences and coordination. In real corridors, releasing vehicles from one intersection can create a platoon that arrives at another intersection later. If the downstream light is not prepared, the released vehicles may create congestion or spillback. A local greedy controller cannot easily anticipate that future arrival if it only reacts to the queue currently visible at each intersection. A reinforcement learning controller can use a fuller state and learn that an upstream action may require a downstream response several steps later.",
    )

    add_heading(doc, "Part II: Coordinated Corridor Extension")
    add_body(
        doc,
        "The extension experiment uses a two-intersection corridor. Intersection A is upstream of Intersection B, and vehicles released from A reach B after a four-step delay. The state includes bucketed queues for both directions at both intersections, the current signal phase at A, the current signal phase at B, an in-transit vehicle bucket, and a demand-regime indicator. The action space contains four joint actions: both signals stay, A stays while B switches, A switches while B stays, or both signals switch. This creates a larger and more coordinated decision problem than the original single-intersection case.",
    )
    add_body(
        doc,
        "The reward is also more system-oriented. It penalizes total queue length, phase switching, and spillback while giving credit for throughput. This is important because raw throughput alone can be misleading. A controller might release many vehicles from A, but if B is not ready to serve them, the network delay can still become worse. The experiment therefore rewards policies that manage the corridor as a whole rather than policies that only clear one visible queue at a time.",
    )
    add_body(
        doc,
        "For the reinforcement learning method, the project uses a Deep Q-Network. A DQN is appropriate because the coordinated corridor has a larger state representation than the original tabular experiment. The model receives the flattened state vector and outputs Q-values for the four joint actions. Training uses experience replay, a target network, and epsilon-greedy exploration. The baselines include static fixed timing, local greedy timing, and a threshold heuristic. Static timing provides a simple control reference. Greedy timing reacts to the currently larger local queue. The threshold heuristic switches only when the opposing queue is larger by a specified margin.",
    )

    add_heading(doc, "Results and Interpretation")
    add_body(
        doc,
        "The complex corridor results support the broader argument of the project. In the coordinated setting, the DQN policy achieved lower average queue and lower average delay than static timing, local greedy timing, and the threshold heuristic. The local greedy baseline still moved many vehicles, but it also switched frequently and produced higher system delay. This pattern is meaningful because it shows the weakness of purely reactive control. Greedy can clear what it sees immediately, but it does not plan for vehicles that are already moving through the corridor and will arrive downstream in the near future.",
    )
    add_body(
        doc,
        "The difference between Part I and Part II is the main lesson. In the simple intersection, greedy remains competitive because the environment is small and the current queue is highly informative. In the coordinated corridor, the current queue is no longer enough. A good policy must understand timing, delayed arrivals, and the relationship between upstream and downstream phases. Reinforcement learning becomes more valuable because it can optimize long-term system behavior instead of only reacting to the largest local queue.",
    )
    add_body(
        doc,
        "This does not mean that reinforcement learning should be described as universally better than heuristic traffic control. Heuristics are often fast, interpretable, and effective in simple settings. The correct conclusion is more specific: reinforcement learning has a structural advantage when the environment contains delayed effects, coordination requirements, and network interactions that are difficult to encode in a simple rule. The extension experiment was designed around exactly those features, so it gives a more convincing reason to continue developing learning-based adaptive traffic control.",
    )

    add_heading(doc, "Conclusion")
    add_body(
        doc,
        "SmartFlow shows a clear progression from a basic MDP demonstration to a richer traffic-control experiment. The original presentation established the core formulation and showed that SARSA could improve on static timing in a single-intersection environment. The extension then showed why the value of reinforcement learning becomes clearer in a coordinated corridor, where decisions have delayed downstream consequences. Future work should expand the environment to larger road networks, continuous state representations, multi-agent signal control, and more realistic safety constraints. The project’s overall conclusion is that reinforcement learning is most promising when traffic control requires coordination over time, not merely reaction to the queue visible at the present moment.",
    )

    doc.add_page_break()
    add_heading(doc, "Appendix", level=1)
    add_appendix_table(
        doc,
        "Table A1. Preserved Presentation Results for Part I",
        ["Approach", "Avg. Reward", "Time"],
        [[r["Approach"], f"{r['Avg Reward']:.2f}", f"{r['Time']:.2f}s"] for _, r in part1.iterrows()],
        [3.0, 1.5, 1.2],
    )
    add_appendix_table(
        doc,
        "Table A2. Complex Corridor Evaluation Results",
        ["Approach", "Reward", "Avg. Queue", "Avg. Delay", "Throughput", "Switches"],
        [
            [
                r["Approach"],
                f"{r['Avg Reward']:.2f}",
                f"{r['Avg Queue']:.2f}",
                f"{r['Avg Delay']:.2f}",
                f"{r['Throughput']:.2f}",
                f"{r['Switches']:.2f}",
            ]
            for _, r in part2.iterrows()
        ],
        [1.25, 1.1, 1.05, 1.05, 1.15, 1.05],
    )
    add_appendix_image(doc, "Figure A1. Part I Reward Comparison", "part1_reward_comparison.png", width=5.8)
    add_appendix_image(doc, "Figure A2. Part I Execution-Time Comparison", "part1_latency_comparison.png", width=5.8)
    add_appendix_image(doc, "Figure A3. Coordinated Corridor Diagram", "part2_network_diagram.png", width=5.8)
    add_appendix_image(doc, "Figure A4. Part II Reward Comparison", "part2_reward_comparison.png", width=5.8)
    add_appendix_image(doc, "Figure A5. Part II Delay Comparison", "part2_delay_comparison.png", width=5.8)
    add_appendix_image(doc, "Figure A6. Part II Queue Comparison", "part2_queue_comparison.png", width=5.8)
    add_appendix_image(doc, "Figure A7. Part II Throughput and Switching Tradeoff", "part2_throughput_switching.png", width=6.2)

    doc.add_page_break()
    add_heading(doc, "References", level=1)
    references = [
        "Mnih, Volodymyr, et al. “Human-Level Control through Deep Reinforcement Learning.” Nature, vol. 518, 2015, pp. 529-533.",
        "Rummery, Gavin A., and Mahesan Niranjan. On-Line Q-Learning Using Connectionist Systems. Cambridge University Engineering Department, 1994.",
        "Sutton, Richard S., and Andrew G. Barto. Reinforcement Learning: An Introduction. 2nd ed., MIT Press, 2018.",
        "Watkins, Christopher J. C. H., and Peter Dayan. “Q-Learning.” Machine Learning, vol. 8, 1992, pp. 279-292.",
        "Wei, Hua, et al. “PressLight: Learning Max Pressure Control to Coordinate Traffic Signals in Arterial Network.” Proceedings of the 25th ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 2019, pp. 1290-1298.",
    ]
    for ref in references:
        add_body(doc, ref)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
